import os
import json
from docx import Document
from docx.shared import RGBColor
from typing import Dict, Any

def get_font_properties(run) -> Dict[str, Any]:
    """
    Extract all font properties from a run with safe attribute checking
    """
    properties = {
        "bold": getattr(run, 'bold', None),
        "italic": getattr(run, 'italic', None),
        "underline": getattr(run, 'underline', None),
        "font_size": str(run.font.size) if hasattr(run.font, 'size') and run.font.size else None,
        "font_name": getattr(run.font, 'name', None),
        "all_caps": getattr(run.font, 'all_caps', None),
        "color": None,
        "highlight_color": str(run.font.highlight_color) if hasattr(run.font, 'highlight_color') else None,
        "subscript": getattr(run.font, 'subscript', None),
        "superscript": getattr(run.font, 'superscript', None),
    }
    
    # Safely extract font color if it exists
    try:
        if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
            rgb = run.font.color.rgb
            properties["color"] = f"#{rgb:06x}" if isinstance(rgb, int) else str(rgb)
    except Exception:
        pass
    
    return properties

def get_paragraph_format(paragraph) -> Dict[str, Any]:
    """
    Extract paragraph formatting properties
    """
    fmt = paragraph.paragraph_format
    return {
        "alignment": str(fmt.alignment) if fmt.alignment else None,
        "first_line_indent": str(fmt.first_line_indent) if fmt.first_line_indent else None,
        "left_indent": str(fmt.left_indent) if fmt.left_indent else None,
        "right_indent": str(fmt.right_indent) if fmt.right_indent else None,
        "line_spacing": str(fmt.line_spacing) if fmt.line_spacing else None,
        "space_before": str(fmt.space_before) if fmt.space_before else None,
        "space_after": str(fmt.space_after) if fmt.space_after else None,
        "keep_together": fmt.keep_together,
        "keep_with_next": fmt.keep_with_next,
        "page_break_before": fmt.page_break_before,
        "widow_control": fmt.widow_control
    }

def extract_docx_content(docx_path: str) -> Dict:
    """
    Extract content from a .docx file and return it as a dictionary with detailed styling
    
    Args:
        docx_path (str): Path to the .docx file
        
    Returns:
        Dict: Dictionary containing document content with detailed styling information
    """
    doc = Document(docx_path)
    
    content = {
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "document_name": os.path.basename(docx_path),
        "lists": []
    }
    
    # Extract section headers with styling
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                if paragraph.text.strip():
                    header_content = {
                        "text": paragraph.text.strip(),
                        "style_name": paragraph.style.name,
                        "formatting": get_paragraph_format(paragraph),
                        "runs": []
                    }
                    
                    for run in paragraph.runs:
                        if run.text.strip():
                            header_content["runs"].append({
                                "text": run.text,
                                "font_properties": get_font_properties(run)
                            })
                    
                    content["headers"].append(header_content)
    
    # Extract paragraphs with detailed styling
    current_list = None
    list_items = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            para_content = {
                "text": para.text.strip(),
                "style_name": para.style.name,
                "formatting": get_paragraph_format(para),
                "runs": [],
                "is_list_item": para.style.name.startswith('List') or hasattr(para, 'numbering'),
                "list_level": para._element.pPr.numPr.ilvl.val if hasattr(para, 'numbering') and hasattr(para._element.pPr, 'numPr') and hasattr(para._element.pPr.numPr, 'ilvl') else None
            }
            
            # Extract styling from individual runs
            for run in para.runs:
                if run.text.strip():
                    para_content["runs"].append({
                        "text": run.text,
                        "font_properties": get_font_properties(run)
                    })
            
            # Handle list items
            if para_content["is_list_item"]:
                if current_list is None:
                    current_list = {"level": para_content["list_level"], "items": []}
                list_items.append(para_content)
            else:
                if current_list is not None:
                    content["lists"].append(list_items)
                    list_items = []
                    current_list = None
                content["paragraphs"].append(para_content)
    
    # Add any remaining list items
    if list_items:
        content["lists"].append(list_items)
    
    # Extract tables with styling
    for table in doc.tables:
        table_data = {
            "style": table.style.name if table.style else "Normal",
            "rows": []
        }
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_content = {
                    "text": cell.text.strip(),
                    "paragraphs": []
                }
                
                # Extract styling from cell paragraphs
                for para in cell.paragraphs:
                    if para.text.strip():
                        para_content = {
                            "text": para.text.strip(),
                            "style_name": para.style.name,
                            "formatting": get_paragraph_format(para),
                            "runs": []
                        }
                        
                        for run in para.runs:
                            if run.text.strip():
                                para_content["runs"].append({
                                    "text": run.text,
                                    "font_properties": get_font_properties(run)
                                })
                                
                        cell_content["paragraphs"].append(para_content)
                
                row_data.append(cell_content)
            table_data["rows"].append(row_data)
        
        if any(any(cell["text"] for cell in row) for row in table_data["rows"]):
            content["tables"].append(table_data)
    
    return content

def process_docx_files():
    """
    Process all .docx files in the current directory and create corresponding JSON files
    """
    docx_files = [f for f in os.listdir('.') if f.lower().endswith('.docx')]
    
    if not docx_files:
        print("No .docx files found in the current directory")
        return
    
    for docx_file in docx_files:
        try:
            print(f"Processing {docx_file}...")
            content = extract_docx_content(docx_file)
            
            json_filename = os.path.splitext(docx_file)[0] + '.json'
            
            with open(json_filename, 'w', encoding='utf-8') as json_file:
                json.dump(content, json_file, ensure_ascii=False, indent=4)
                
            print(f"Successfully created {json_filename}")
            
        except Exception as e:
            print(f"Error processing {docx_file}: {str(e)}")
            print(f"Full error details: ", e)

if __name__ == "__main__":
    process_docx_files()


